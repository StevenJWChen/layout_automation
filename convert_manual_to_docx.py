#!/usr/bin/env python3
"""
Convert GDS User Manual from Markdown to DOCX
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def markdown_to_docx(md_file: str, docx_file: str):
    """Convert markdown file to DOCX with professional styling"""

    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Track state
    in_code_block = False
    code_buffer = []
    in_table = False
    table_lines = []

    for line in lines:
        line = line.rstrip()

        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_buffer)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                # Add shading
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F8F8F8')
                p._p.get_or_add_pPr().append(shading_elm)
                code_buffer = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        # Handle headings
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            continue

        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            p.runs[0].font.color.rgb = RGBColor(0, 136, 204)
            continue

        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            p.runs[0].font.color.rgb = RGBColor(0, 170, 204)
            continue

        elif line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)
            continue

        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph()
            continue

        # Handle lists
        elif re.match(r'^[\*\-]\s+', line):
            text = re.sub(r'^[\*\-]\s+', '', line)
            text = process_inline_markdown(text)
            p = doc.add_paragraph(text, style='List Bullet')
            continue

        elif re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            text = process_inline_markdown(text)
            p = doc.add_paragraph(text, style='List Number')
            continue

        # Handle tables
        elif '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            continue

        else:
            # If we were in a table, process it
            if in_table:
                create_table(doc, table_lines)
                table_lines = []
                in_table = False

            # Regular paragraph
            if line.strip():
                text = process_inline_markdown(line)
                p = doc.add_paragraph(text)
            else:
                doc.add_paragraph()

    # Process any remaining table
    if in_table and table_lines:
        create_table(doc, table_lines)

    # Save document
    print(f"Converting {md_file} to DOCX...")
    doc.save(docx_file)
    print(f"✓ Created: {docx_file}")

    # Get file size
    size_kb = Path(docx_file).stat().st_size / 1024
    print(f"  File size: {size_kb:.1f} KB")

def process_inline_markdown(text: str) -> str:
    """Process inline markdown like bold, italic, code"""
    # Note: python-docx doesn't support rich text in paragraph text
    # This is a simplified version that removes markdown syntax
    # For proper formatting, we'd need to use runs

    # Remove code backticks for now
    text = re.sub(r'`([^`]+)`', r'\1', text)

    # Remove bold
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)

    # Remove italic
    text = re.sub(r'\*([^*]+)\*', r'\1', text)

    return text

def create_table(doc, table_lines):
    """Create a table from markdown table lines"""
    # Parse table
    rows = []
    for line in table_lines:
        if line.strip().startswith('|---') or line.strip().startswith('| ---'):
            continue  # Skip separator line
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)

    if not rows:
        return

    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'

    # Fill table
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = process_inline_markdown(cell_text)

            # Header row formatting
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                # Add blue background to header
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '0066CC')
                cell._tc.get_or_add_tcPr().append(shading_elm)

if __name__ == '__main__':
    markdown_to_docx('GDS_USER_MANUAL.md', 'GDS_USER_MANUAL.docx')
